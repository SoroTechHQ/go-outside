#!/usr/bin/env python3
"""Generate GoOutside Mini Brand Guide as a .docx file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "../documents/GoOutside_Brand_Guide_TEMP.docx")

def hex_to_rgb(hex_color: str):
    """Convert #rrggbb to (r, g, b) tuple."""
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex="#0f110f"):
    h = doc.add_heading(text, level=level)
    h.clear()
    run = h.add_run(text)
    r, g, b = hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.bold = True
    run.font.size = Pt({1: 22, 2: 16, 3: 13}[level])
    return h

def add_label(doc, text, color_hex="#6f6f6f"):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    r, g, b = hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.size = Pt(8)
    run.font.bold = True
    return p

def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x0f, 0x11, 0x0f)
    return p

def color_swatch_row(doc, swatches):
    """
    swatches = list of (label, hex, description) tuples.
    Renders a table row with a color block + label + hex + description.
    """
    table = doc.add_table(rows=len(swatches), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for i, (label, hex_color, description) in enumerate(swatches):
        row = table.rows[i]

        # Col 0: Color swatch
        swatch_cell = row.cells[0]
        set_cell_bg(swatch_cell, hex_color)
        swatch_cell.width = Cm(2)
        swatch_cell.paragraphs[0].add_run("   ")
        p = swatch_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # Col 1: Label + hex
        name_cell = row.cells[1]
        name_cell.width = Cm(5)
        p1 = name_cell.paragraphs[0]
        r = p1.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        p2 = name_cell.add_paragraph()
        r2 = p2.add_run(hex_color.upper())
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0x6f, 0x6f, 0x6f)

        # Try to calculate RGB string
        try:
            rgb = hex_to_rgb(hex_color)
            rgb_str = f"RGB({rgb[0]}, {rgb[1]}, {rgb[2]})"
            p3 = name_cell.add_paragraph()
            r3 = p3.add_run(rgb_str)
            r3.font.size = Pt(8)
            r3.font.color.rgb = RGBColor(0xa9, 0xa9, 0xa9)
        except Exception:
            pass

        # Col 2: Description
        desc_cell = row.cells[2]
        desc_cell.width = Cm(9)
        p_desc = desc_cell.paragraphs[0]
        r_desc = p_desc.add_run(description)
        r_desc.font.size = Pt(9)
        r_desc.font.color.rgb = RGBColor(0x6f, 0x6f, 0x6f)

    doc.add_paragraph()

def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xd8, 0xd8, 0xd8)

# ─────────────────────────────────────────────────────────────────────────────
# Build Document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# ── Cover ─────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("GoOutside")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0x2f, 0x8f, 0x45)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("TEMPORARY MINI BRAND GUIDE")
r2.font.size = Pt(11)
r2.font.bold = True
r2.font.color.rgb = RGBColor(0xa9, 0xa9, 0xa9)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Current state audit — June 2026\nFor internal use. Placeholder for full rebrand.")
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0xa9, 0xa9, 0xa9)
r3.font.italic = True

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 1. BRAND IDENTITY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. Brand Identity", 1, "#0f110f")

add_heading(doc, "App Name & Tagline", 2)
add_label(doc, "Primary Name")
add_body(doc, "GoOutside")
add_label(doc, "Short Tagline")
add_body(doc, "What's on in Accra?")
add_label(doc, "Full Tagline")
add_body(doc, "Social-first event discovery for Ghana. Find events, earn Pulse Points, go out.")
add_label(doc, "Domain / URL")
add_body(doc, "gooutside.club")
add_label(doc, "Locale")
add_body(doc, "English (Ghana) — en_GH")
doc.add_paragraph()

add_heading(doc, "Brand Essence", 2)
add_body(doc, (
    "GoOutside is a social event discovery platform built specifically for Ghana. "
    "It blends local culture, social proof, and gamified loyalty (Pulse Points) to "
    "help people find and attend events — and reward them for doing so. "
    "The brand is energetic, community-first, and distinctly Ghanaian."
))
doc.add_paragraph()

add_heading(doc, "Brand Voice & Tone", 2)
add_label(doc, "Personality")
add_body(doc, "Energetic · Local · Social · Clean · Trustworthy")
add_label(doc, "Writing Style")
add_body(doc,
    "• Short, punchy copy. Action-oriented headlines.\n"
    "• Eyebrow labels in ALL CAPS with wide letter-spacing (e.g. 'FRIDAY IN OSU', 'BUILDERS AND OPERATORS').\n"
    "• Feature sections open with a branded mini-label in green, then a larger serif/italic h2.\n"
    "• Numbers and social proof are highlighted (e.g. '47 friends going', '68 left').\n"
    "• Ghanaian place names and cultural references used prominently (Osu, Legon, East Legon, Accra, etc.).\n"
    "• Avoid emojis — use Phosphor icons instead."
)
add_label(doc, "Example Headlines Found in App")
add_body(doc,
    "• \"What's on in Accra?\"\n"
    "• \"Ga Rooftop After Hours\" (eyebrow: FRIDAY IN OSU)\n"
    "• \"Discover\" (section eyebrow in green)\n"
    "• \"GoOutside gives you everything you need to sell tickets, manage attendees, and grow your audience in Ghana.\"\n"
    "• \"I never knew there were this many things happening in Accra on weekends.\""
)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 2. LOGO & ASSETS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. Logo & Visual Assets", 1, "#0f110f")

add_heading(doc, "Logo Files (current)", 2)
add_label(doc, "Web App")
add_body(doc,
    "• /public/logo-full.png  — Full wordmark (add image here)\n"
    "• /public/logo-mini.png  — Icon / compact mark (add image here)\n"
    "• /public/favicon-icon.png  — Favicon\n"
    "• /app/icon.png  — App icon (used in browser tab)"
)
add_label(doc, "Admin App")
add_body(doc,
    "• /public/logo-full.png  — same full wordmark\n"
    "• /public/logo-mini.png  — same icon\n"
    "• /public/favicon-icon.png  — Favicon"
)
add_label(doc, "Figma Exports Available")
add_body(doc,
    "• documents/ticket-inspo/figma-exports/Main App.svg\n"
    "• documents/ticket-inspo/figma-exports/GOLD TICKETS.svg\n"
    "• documents/ticket-inspo/figma-exports/SILVER TICKETS.svg\n"
    "• documents/ticket-inspo/figma-exports/DEFAULT TICKETS 2.svg\n"
    "• documents/ticket-inspo/figma-exports/CARDS for tickets.svg"
)

add_body(doc, "\n[  PASTE LOGO IMAGES HERE  ]\n", 10)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 3. COLOR SYSTEM — WEB APP
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. Color System — Web App (apps/web)", 1, "#0f110f")

# Primary
add_heading(doc, "3.1 Primary Brand Color — Light Mode", 2)
color_swatch_row(doc, [
    ("Brand Green (Primary)", "#2f8f45", "Primary CTA buttons, active states, links, icons, send buttons, badges. Used everywhere a 'brand' moment is needed."),
    ("Brand Hover", "#256f36", "Hover state for Brand Green buttons."),
    ("Brand Dim / Glow", "#2f8f45", "rgba(47,143,69,0.10) — subtle green tint on hover/active surfaces, avatar backgrounds, highlight backgrounds."),
])

add_heading(doc, "3.2 Primary Brand Color — Dark Mode", 2)
color_swatch_row(doc, [
    ("Brand Green (Dark)", "#087f18", "var(--brand) in dark mode. Slightly deeper/richer forest green."),
    ("Brand Hover (Dark)", "#2ec962", "Hover glow in dark. Lighter lime-green to pop on dark backgrounds."),
    ("Chat Green (Accent)", "#5FBF2A", "Used in Stream Chat send bubble, organizer verified badge, message send button. Brighter lime-green variant."),
])

# Secondary / Accent
add_heading(doc, "3.3 Secondary & Accent Colors", 2)
color_swatch_row(doc, [
    ("Pink / Danger", "#e85d8a", "Danger states, decline buttons, error borders. Used for status-review-text, danger-dim, hover on decline actions."),
    ("Blue / Link", "#4a7ae8", "Secondary accent. Used for links, blue-themed UI elements."),
    ("Unread Dot Blue", "#0095F6", "Unread message indicator dot in DM sidebar."),
    ("Online Status Green", "#00BA88", "Online presence dot on user avatars in chat."),
    ("Amber Warning", "#f59e0b", "Message request / pending state label."),
])

# Pulse Points Gold
add_heading(doc, "3.4 Pulse Points Gold (Loyalty System)", 2)
color_swatch_row(doc, [
    ("Pulse Gold", "#bf9150", "var(--pulse-gold). Used for Pulse Points branding, gold tier elements, wallet accents."),
    ("Pulse Gold Soft", "#bf9150", "rgba(191,145,80,0.12) — soft background tint for Pulse Gold areas."),
    ("Pulse Gold Border", "#bf9150", "rgba(191,145,80,0.24) — border for Pulse Gold cards/chips."),
])

# Tier Colors
add_heading(doc, "3.5 Pulse Score Tier Colors", 2)
color_swatch_row(doc, [
    ("Legend (2000+ pts)", "#DAA520",  "Goldenrod — top tier."),
    ("City Native (1000–1999)", "#c87c2a", "Burnt amber."),
    ("Scene Kid (600–999)", "#4a9f63",  "Mid green."),
    ("Regular (300–599)", "#4a9f63",   "Mid green."),
    ("Explorer (100–299)", "#4a9f63",  "Mid green."),
    ("Newcomer (0–99)", "#9CA3AF",    "Cool grey — neutral entry tier."),
])

# Category Colors
add_heading(doc, "3.6 Event Category Colors", 2)
color_swatch_row(doc, [
    ("Music", "#7c3aed",  "Purple — used on category badges and pills."),
    ("Tech",  "#2563eb",  "Blue."),
    ("Food",  "#d97706",  "Amber."),
    ("Arts",  "#be185d",  "Hot pink."),
    ("Sports","#059669",  "Emerald green."),
    ("Networking","#0891b2","Cyan."),
    ("Education","#b45309","Brown amber."),
    ("Community","#5b21b6","Dark violet."),
])

# Backgrounds
add_heading(doc, "3.7 Background & Surface Colors — Light Mode", 2)
color_swatch_row(doc, [
    ("Page / App Base", "#ffffff",   "var(--bg-base, --bg-app). Pure white."),
    ("Surface / Elevated", "#f7f7f7","Cards, sidebars, elevated containers."),
    ("Card", "#ffffff",              "Default card background."),
    ("Card Alt", "#f1f1f1",          "Secondary card variant."),
    ("Glass", "#ffffff",             "rgba(255,255,255,0.88) — glass/blur overlay."),
    ("Muted", "#f3f3f3",             "Muted backgrounds, chip surfaces."),
    ("Onboarding BG", "#f8f8f8",     "Onboarding flow page background."),
])

add_heading(doc, "3.8 Background & Surface Colors — Dark Mode", 2)
color_swatch_row(doc, [
    ("Page / App Base (Dark)", "#0c0f0d",  "Near-black with green undertone."),
    ("Surface (Dark)", "#111411",           "Slightly lighter surface."),
    ("Card (Dark)", "#161a16",              "Card background."),
    ("Card Alt (Dark)", "#1d231e",          "Secondary card."),
    ("Glass (Dark)", "#161a16",             "rgba(22,26,22,0.86)."),
])

# Text
add_heading(doc, "3.9 Text Colors", 2)
color_swatch_row(doc, [
    ("Text Primary (Light)", "#0f110f",   "Near-black with a slight green tint. Used for all main body and heading text."),
    ("Text Secondary (Light)", "#6f6f6f", "Subtext, metadata, captions."),
    ("Text Tertiary (Light)", "#a9a9a9",  "Placeholder text, faint labels."),
    ("Text Primary (Dark)", "#edf1ee",    "Off-white with a slight green tint."),
    ("Text Secondary (Dark)", "#7a9080",  "Muted green-grey."),
    ("Text Tertiary (Dark)", "#3d5240",   "Very muted for de-emphasised content."),
])

# Status
add_heading(doc, "3.10 Status / Semantic Colors", 2)
color_swatch_row(doc, [
    ("Live (Light)", "#2f8f45",  "Status badge: live events. Same as brand green."),
    ("Pending (Light)", "#7a4800", "Amber-brown for pending states."),
    ("Review (Light)", "#7a1840", "Dark pink for review/warning states."),
    ("Live (Dark)", "#45b423",   "Slightly lighter green for dark backgrounds."),
    ("Error / Danger", "#e85d8a","Pink-red for errors and danger actions."),
    ("Error Text", "#f87171",    "Tailwind red-400 for error messages."),
])

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 4. COLOR SYSTEM — ADMIN APP
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. Color System — Admin App (apps/admin)", 1, "#0f110f")
add_body(doc, "The admin app uses a slightly different green variant and adds a richer accent palette.")
doc.add_paragraph()

add_heading(doc, "4.1 Primary Color", 2)
color_swatch_row(doc, [
    ("Brand (Admin Light)", "#1fa764", "Slightly warmer mid-green. Used the same way as #2f8f45 in the web app."),
    ("Brand (Admin Dark)", "#3ddc97",  "Bright mint/teal-green for dark mode."),
])

add_heading(doc, "4.2 Admin Accent Palette", 2)
add_body(doc, "The admin app exposes a richer set of named accent colors not present in the consumer web app:")
color_swatch_row(doc, [
    ("Accent Cyan",   "#38bdf8", "Used for secondary highlights, data viz."),
    ("Accent Violet", "#a78bfa", "Used for tagging, badges."),
    ("Accent Coral",  "#fb7185", "Danger / error states (alias: --pink)."),
    ("Accent Amber",  "#f59e0b", "Warnings, pending states."),
    ("Accent Lime",   "#84cc16", "Positive trends, growth indicators."),
    ("Admin Blue",    "#4f7cff", "Link/action blue."),
])

add_heading(doc, "4.3 Admin Backgrounds", 2)
color_swatch_row(doc, [
    ("Base (Admin Light)", "#f6f8f6", "Very light green-tinted white."),
    ("Elevated", "#ffffff",           "White elevated cards."),
    ("Card Alt", "#eef4f2",           "Green-tinted alt card."),
    ("Base (Admin Dark)", "#09110d",  "Deep forest near-black."),
    ("Elevated (Dark)", "#0f1914",    "Slightly lighter forest."),
    ("Card (Dark)", "#131d18",        "Dark green card background."),
])

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 5. TYPOGRAPHY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. Typography", 1, "#0f110f")

add_heading(doc, "5.1 Web App Fonts", 2)
add_label(doc, "Body / UI Font")
add_body(doc,
    "Inter\n"
    "Weights: 300 (Light), 400 (Regular), 500 (Medium), 600 (SemiBold), 700 (Bold)\n"
    "CSS variable: --font-body\n"
    "Usage: All body text, UI labels, buttons, inputs, navigation, everything by default."
)
add_label(doc, "Display / Heading Font")
add_body(doc,
    "Inter (same family, different variable)\n"
    "Weights: 400, 500, 600, 700\n"
    "CSS variable: --font-display\n"
    "Usage: Large headings, hero titles. Same typeface used — distinction is by weight and size.\n"
    "Note: Both fonts are Inter; the codebase reserves --font-display for future replacement with a distinct display typeface."
)

add_heading(doc, "5.2 Admin App Fonts", 2)
add_label(doc, "Body Font")
add_body(doc,
    "DM Sans\n"
    "Weights: 300, 400, 500, 600, 700\n"
    "CSS variable: --font-body\n"
    "Usage: All admin UI body text, tables, sidebar items."
)
add_label(doc, "Display Font")
add_body(doc,
    "Inter\n"
    "Weights: 400, 500, 600, 700\n"
    "CSS variable: --font-display\n"
    "Usage: Admin headings, data labels."
)

add_heading(doc, "5.3 Type Scale — Web App (found in use)", 2)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "Size / Class"
hdr[2].text = "Usage"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

rows_data = [
    ("Hero H1 (Landing)", "40–72px / font-extrabold, tracking-tight", "Main landing page hero headline. Scales from 40px mobile → 72px desktop."),
    ("Section H2", "28–36px / font-bold or font-normal italic", "Feature section headings. Sometimes italic for editorial feel."),
    ("Card Title", "18px / font-bold", "Event card titles, profile names on cards."),
    ("Body / Description", "15–18px / font-light", "Subheadings, descriptions below hero."),
    ("UI Body", "14–15px / font-normal 400", "General UI text, feed items, chat previews."),
    ("Caption / Meta", "12–13px / font-medium or 400", "Dates, locations, timestamps, sub-labels."),
    ("Eyebrow Label", "9–11px / font-bold uppercase tracking-[0.14–0.18em]", "Section eyebrows above headings (e.g. 'FRIDAY IN OSU'). Brand green or muted."),
    ("Badge / Tag", "9–10px / font-bold uppercase tracking-[0.08em]", "Category pills, event type badges, status tags."),
    ("Sidebar Nav", "16px / font-semibold", "Sidebar navigation items."),
    ("Progress Step", "varies / font-semibold", "Onboarding step indicators."),
]
for label, scale, usage in rows_data:
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = scale
    row.cells[2].text = usage
    for i, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

add_heading(doc, "5.4 Type Conventions", 2)
add_body(doc,
    "• Eyebrow labels above headings always use ALL CAPS + wide letter spacing (tracking-[0.14em] to tracking-[0.18em]).\n"
    "• Brand green (#2f8f45) is used for eyebrow labels to signal section starts.\n"
    "• Hero headlines use font-extrabold (800+) with tracking-tight.\n"
    "• Chat UI uses 0.875rem (14px) for messages, 0.96rem (15px) for compose input.\n"
    "• Notification badges use font-bold, 9px, white on brand-green background.\n"
    "• The .font-display class applies var(--font-display); .font-body applies var(--font-body)."
)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 6. BORDER RADIUS / SHAPE SYSTEM
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Shape & Border Radius System", 1, "#0f110f")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Token"
hdr[1].text = "Value"
hdr[2].text = "Usage"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

radius_data = [
    ("--radius-pill", "9999px", "Buttons, badges, chips, pills, avatar borders, tab active states. Everything rounded-full."),
    ("--radius-card-lg", "20px", "Large cards, main content cards on feed. Tailwind: rounded-card."),
    ("--radius-card", "14px",   "Standard inner cards, event cards, post cards."),
    ("--radius-panel", "12px",  "Panels, modals, dropdowns, side panes."),
    ("--radius-input", "8px",   "Form inputs, text areas (non-pill)."),
    ("--radius-image", "8px",   "Event/post images inside cards."),
    ("--radius-tag", "7px",     "Small label tags."),
    ("card (Tailwind)", "20px",  "rounded-card in tailwind.config.ts."),
    ("card-lg (Tailwind)", "32px","rounded-card-lg — extra large cards."),
    ("panel (Tailwind)", "16px", "rounded-panel — panel variant."),
]
for token, value, usage in radius_data:
    row = table.add_row()
    row.cells[0].text = token
    row.cells[1].text = value
    row.cells[2].text = usage
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_body(doc, "Key pattern: The UI is predominantly pill-shaped for interactive elements (buttons, chips, badges, inputs) and uses 14–20px radius for card containers. Very little sharp/square elements remain.")
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 7. ICONOGRAPHY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. Iconography", 1, "#0f110f")
add_label(doc, "Icon Library")
add_body(doc,
    "Primary: @phosphor-icons/react\n"
    "Secondary: lucide-react (used in some admin and shared UI components)\n"
    "Radix UI icons (used within Radix UI primitives)\n\n"
    "Rule: No emojis in UI — always substitute with Phosphor icons.\n"
    "Icon sizes commonly used: 16px, 18px, 20px, 24px.\n"
    "Icons are typically colored with var(--text-secondary) at rest, var(--brand) or white on active/hover."
)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 8. SHADOWS & ELEVATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "8. Shadows & Elevation", 1, "#0f110f")

table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Token / Class"
hdr[1].text = "Value & Usage"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.size = Pt(9)

shadow_data = [
    ("--card-shadow", "0 1px 3px rgba(0,0,0,0.06), 0 1px 2px rgba(0,0,0,0.04)\nDefault card shadow. Very subtle."),
    ("--card-shadow-hover", "0 8px 24px rgba(0,0,0,0.10), 0 2px 6px rgba(0,0,0,0.06)\nCard hover / lift state."),
    ("--home-shadow", "0 1px 3px rgba(0,0,0,0.06), 0 1px 2px rgba(0,0,0,0.04)\nHome feed cards."),
    ("--home-shadow-strong", "0 12px 30px rgba(0,0,0,0.08), 0 2px 6px rgba(0,0,0,0.05)\nSide pane, expanded cards."),
    ("--home-active-shadow", "0 0 0 1px rgba(47,143,69,0.16), 0 8px 30px rgba(0,0,0,0.05)\nActive/selected card with green ring."),
    ("--brand-shadow", "0 12px 30px rgba(15,17,15,0.08)\nBrand-green shadow on CTAs."),
    ("Button CTA shadow", "0 2px 12px rgba(47,143,69,0.40)\nGreen glow on primary buttons (landing page)."),
    ("Compose card shadow", "0 28px 80px rgba(0,0,0,0.14)\nMessage compose modal."),
    ("glass-card class", "var(--home-shadow-strong) + backdrop-filter: blur(24px)\nGlass morphism cards (hero overlay, side panels)."),
]
for token, value in shadow_data:
    row = table.add_row()
    row.cells[0].text = token
    row.cells[1].text = value
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 9. COMPONENT PATTERNS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "9. Key Component Patterns", 1, "#0f110f")

add_heading(doc, "Buttons", 2)
add_body(doc,
    "PRIMARY — rounded-full, bg-brand (#2f8f45), white text, font-bold, shadow with green glow.\n"
    "  Example: bg-[#2f8f45] px-5 h-10 rounded-full text-[14px] font-bold text-white shadow-[0_2px_12px_rgba(47,143,69,0.40)]\n\n"
    "SECONDARY OUTLINE — rounded-full, transparent bg, brand-green border and text.\n"
    "  Example: border border-[rgba(47,143,69,0.35)] px-7 text-[14px] font-bold text-[#2f8f45]\n\n"
    "GHOST / ICON — no bg, no border, hover adds bg-surface.\n\n"
    "DANGER — bg-pink (#e85d8a) or brand-dim bg with pink text."
)

add_heading(doc, "Cards", 2)
add_body(doc,
    "Standard event cards: rounded-[14px] or rounded-[20px], border border-subtle, bg-card.\n"
    "Large hero cards: rounded-[20px] with image scrim overlay, white foreground text.\n"
    "Glass cards: glass-card class (bg-glass + backdrop-blur-[24px] + border-subtle).\n"
    "Hover state: border-color shifts to brand green, shadow lifts."
)

add_heading(doc, "Badges / Pills / Chips", 2)
add_body(doc,
    "All use rounded-full (pill shape).\n"
    "Category badges: colored per CATEGORY_COLORS map above.\n"
    "Status 'LIVE': bg green-tinted, border green-tinted, text brand green.\n"
    "Pulse tier pill: tier color, semi-transparent bg tint.\n"
    "Eyebrow label chips: transparent, brand-green border at 50% opacity, tracking-[0.14em] uppercase.\n"
    "Notification count: bg-brand, white text, font-bold, 9px."
)

add_heading(doc, "Inputs", 2)
add_body(doc,
    "Default: rounded-[8px] or rounded-full, border-subtle, bg-card.\n"
    "Focus: border shifts to brand-green at 55% opacity + green box-shadow ring.\n"
    "Chat compose: rounded-[22px], focus border rgba(95,191,42,0.5).\n"
    "Onboarding inputs: bg-[#f5f5f5] (light), rounded, border-[rgba(0,0,0,0.10)], focus ring brand green."
)

add_heading(doc, "Navigation", 2)
add_body(doc,
    "Sidebar (desktop): 72px wide collapsed, fixed left. bg-elevated. Nav items use Phosphor icons + labels.\n"
    "Active nav item: brand green icon, green dot indicator, brand-green ring on avatar.\n"
    "Bottom nav (mobile): fixed bottom, bg-card, blur backdrop, pill-shaped active indicator.\n"
    "Top loader bar: brand green (#2f8f45), 3px height."
)

add_heading(doc, "Chat (Stream Chat)", 2)
add_body(doc,
    "Sent messages: bg #5FBF2A (lime green), white text, border-radius 18px 18px 5px 18px.\n"
    "Received messages: bg #F0F0F0 (light grey), dark text, border-radius 18px 18px 18px 5px.\n"
    "Send button: bg-brand, rounded-full, white icon.\n"
    "Organizer thread accent: left border rgba(95,191,42,0.45).\n"
    "Unread dot: #0095F6 (blue).\n"
    "Online status: #00BA88 (teal green)."
)

add_heading(doc, "Background Texture", 2)
add_body(doc,
    "Page body gradient: radial green glow at top-left (4% opacity), secondary green accent at top-right, then linear fade to bg-app.\n"
    "Grid overlay: linear + cross-hatch grid lines at rgba(0,0,0,0.03–0.035), fading out with a mask. Opacity 0.14. Decorative only.\n"
    "Glow orbs: animated radial blobs (glow-orb class, 6–8s ease-in-out pulse).\n"
    "Glass morphism: backdrop-filter blur(24px) on side panels and compose modals."
)

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 10. MOTION / ANIMATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "10. Motion & Animation", 1, "#0f110f")
add_body(doc,
    "Library: framer-motion (Framer Motion) for page transitions and component animations.\n"
    "Library: atropos for 3D tilt cards (used on ticket cards in wallet).\n\n"
    "Standard transitions: 0.15s–0.20s ease for color, background, border-color changes.\n"
    "Glow orb animation: 6–8s ease-in-out infinite pulse (opacity 0.06 → 0.10 + scale 1 → 1.05).\n"
    "QR sweep bar: 3s ease-in-out infinite horizontal sweep with green gradient.\n"
    "Skeleton shimmer: translateX(-100%) → translateX(100%) for loading placeholders.\n"
    "Top loader: easing 'ease', 300ms speed, brand green.\n"
    "Button hover: transform scale(1.04) on send button hover.\n"
    "Sidebar open/close: 250–280ms cubic-bezier(0.4,0,0.2,1) / (0.22,1,0.36,1)."
)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 11. DARK MODE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "11. Dark Mode Behaviour", 1, "#0f110f")
add_body(doc,
    "Trigger: :root[data-theme='dark'] selector (not prefers-color-scheme).\n"
    "Set by: ThemeScript component from @gooutside/ui, stored in localStorage or system pref.\n\n"
    "Key dark mode shifts:\n"
    "• Brand green shifts from #2f8f45 → #087f18 (web) or #1fa764 → #3ddc97 (admin).\n"
    "• Chat green stays #5FBF2A for send bubble consistency.\n"
    "• Backgrounds shift to near-black with green undertones (#0c0f0d, #111411, #161a16).\n"
    "• Text goes from #0f110f → #edf1ee (slight warm-green tint white).\n"
    "• Borders become white-based at 6–12% opacity.\n"
    "• Glow effects intensify (0.18 opacity vs 0.10 in light).\n"
    "• Admin dark base is deeper forest green (#09110d vs #0c0f0d)."
)
doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 12. PAGES — SURFACE NOTES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "12. Per-Surface Color Notes", 1, "#0f110f")

add_heading(doc, "Landing Page (/)", 2)
add_body(doc,
    "Dark-toned hero with full-bleed event image. White foreground text.\n"
    "Hero pill badge: green border (rgba(47,143,69,0.50)), dark translucent bg, #86efac text.\n"
    "CTA button: solid brand green + green glow shadow.\n"
    "Secondary CTA: ghost green outline.\n"
    "Social proof stats: #4ade80 (bright green) numbers on dark bg.\n"
    "Feature sections (below fold): white background, eyebrow in #2f8f45, h2 in #0f110f.\n"
    "Testimonial cards: white bg, standard card styling.\n"
    "Footer: dark bg with logo."
)

add_heading(doc, "Home Feed (/home)", 2)
add_body(doc,
    "Light mode by default. Subtle radial green glow behind header.\n"
    "Category rail: rounded-full chip pills, active chip = brand green bg, white text.\n"
    "Event cards: white bg, rounded-[14px], border-subtle.\n"
    "Hero carousel: full-bleed images with dark gradient scrim.\n"
    "Section headers: bold labels with Phosphor icons."
)

add_heading(doc, "Onboarding (/onboarding/*)", 2)
add_body(doc,
    "Clean, minimal. Light: #f8f8f8 bg, white card, brand-green progress bar and selected chips.\n"
    "Dark: near-black (#0a0a0a) bg, #161616 card.\n"
    "Progress bar: brand green fill.\n"
    "Selected interest chips: brand-green border + bg-tint + brand-green text."
)

add_heading(doc, "Dashboard (all /dashboard/* pages)", 2)
add_body(doc,
    "Inherits page-grid layout with grid-line texture overlay.\n"
    "Cards use var(--bg-card) + var(--border-subtle).\n"
    "Active states use brand green.\n"
    "Rewards page: Pulse Gold accents for PP balances, Gold tier elements.\n"
    "Messages: Stream Chat with brand-green outgoing bubbles."
)

add_heading(doc, "Organizer Dashboard (/organizer/*)", 2)
add_body(doc,
    "Admin-style layout. Sidebar with DM Sans (body), Inter (headings).\n"
    "Admin app uses the brighter #1fa764 (light) / #3ddc97 (dark) green variant.\n"
    "Additional accent colors (cyan, violet, coral, amber, lime) used in charts and analytics.\n"
    "Status chips follow same pill-shape pattern."
)

add_heading(doc, "Search (/search)", 2)
add_body(doc,
    "AI Chat Panel: dark-toned panel, brand-green user message bubbles.\n"
    "Search bar: rounded-full, expands on focus with brand-green ring.\n"
    "Result pills: category-colored badges."
)

add_heading(doc, "Tickets & Wallet", 2)
add_body(doc,
    "Ticket cards use Atropos 3D tilt effect.\n"
    "QR code screen: brand-green sweep bar animation.\n"
    "Wallet: Pulse Gold (#bf9150) for PP balance display.\n"
    "Ticket tiers: Gold (#bf9150), Silver (#9CA3AF) color themes per Figma exports."
)

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# 13. NOTES & GAPS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "13. Notes, Inconsistencies & Gaps", 1, "#0f110f")
add_body(doc,
    "The following are observations from the current codebase audit:\n\n"
    "1. TWO GREEN VARIANTS COEXIST:\n"
    "   Web app primary: #2f8f45 (darker forest green)\n"
    "   Chat send bubble / verified badge: #5FBF2A (brighter lime-green)\n"
    "   Admin light: #1fa764 (warmer mid-green)\n"
    "   → A unified single brand green should be chosen for the rebrand.\n\n"
    "2. FONT DECISION:\n"
    "   Web uses Inter for both body and display (same family).\n"
    "   Admin uses DM Sans (body) + Inter (display).\n"
    "   → Consider a distinct display typeface for GoOutside hero text.\n\n"
    "3. RADIUS INCONSISTENCY:\n"
    "   Buttons are always pill (9999px). Cards range from 14px to 32px.\n"
    "   Some hardcoded values in JSX, some via CSS tokens.\n"
    "   → Consolidate all radius values into CSS token system.\n\n"
    "4. ADMIN VS WEB PALETTES:\n"
    "   Admin has extra accent colors (cyan, violet, amber, lime) not used in web app.\n"
    "   → Decide if these accents should be brought into the web app as part of rebrand.\n\n"
    "5. DARK MODE BRAND COLOR INCONSISTENCY:\n"
    "   Dark mode brand: #087f18 (web) — very dark, may be too low contrast.\n"
    "   Admin dark: #3ddc97 — much brighter mint, higher contrast.\n"
    "   → Align dark mode brand green across both apps.\n\n"
    "6. LOGO FILES:\n"
    "   PNGs available but no SVG wordmark in /public.\n"
    "   Figma SVGs exist in documents/ticket-inspo/figma-exports/.\n"
    "   → Export clean SVG versions of both logo-mini and logo-full.\n\n"
    "7. OG IMAGE:\n"
    "   Generated dynamically at /api/og — no static brand card on file.\n"
    "   → Create a static OG image template as part of rebrand.\n\n"
    "8. NO CUSTOM DISPLAY FONT YET:\n"
    "   Both --font-body and --font-display resolve to Inter on the web app.\n"
    "   --font-display is an alias ready to be pointed at a different font.\n"
)

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run("GoOutside Mini Brand Guide — Temporary Audit Document\nJune 2026 · For internal use")
r_end.font.size = Pt(9)
r_end.font.color.rgb = RGBColor(0xa9, 0xa9, 0xa9)
r_end.font.italic = True

# ─────────────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
